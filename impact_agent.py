#!/usr/bin/env python3
"""
Impact Analysis Agent - Render Optimized
"""

from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import json
import os
import urllib.request
import urllib.error
import subprocess
import tempfile
import shutil
from datetime import datetime
from io import BytesIO
import uvicorn
import time

# Load environment variables
def load_env():
    try:
        with open('.env', 'r') as f:
            for line in f:
                if line.strip() and '=' in line and not line.startswith('#'):
                    key, value = line.split('=', 1)
                    os.environ[key.strip()] = value.strip()
    except:
        pass

load_env()

try:
    import PyPDF2
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    PDF_SUPPORT = True
except:
    PDF_SUPPORT = False

try:
    from docx import Document
    DOC_SUPPORT = True
except:
    DOC_SUPPORT = False

# Configuration
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "").strip()
GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.1-8b-instant").strip()
GROQ_TEMPERATURE = float(os.getenv("GROQ_TEMPERATURE", "0.3"))
GROQ_URL = os.getenv("GROQ_URL", "https://api.groq.com/openai/v1/chat/completions").strip()

# Configurable limits and timeouts
GIT_CLONE_TIMEOUT = int(os.getenv("GIT_CLONE_TIMEOUT", "60"))
API_REQUEST_TIMEOUT = int(os.getenv("API_REQUEST_TIMEOUT", "30"))
MAX_TOKENS = int(os.getenv("MAX_TOKENS", "4500"))
FILE_CONTENT_LIMIT = int(os.getenv("FILE_CONTENT_LIMIT", "2000"))
ARCH_CONTENT_LIMIT = int(os.getenv("ARCH_CONTENT_LIMIT", "1500"))
KEY_FILE_CONTENT_LIMIT = int(os.getenv("KEY_FILE_CONTENT_LIMIT", "800"))
MAX_KEY_FILES = int(os.getenv("MAX_KEY_FILES", "10"))
MAX_LANGUAGES = int(os.getenv("MAX_LANGUAGES", "5"))
MAX_MOCK_FILES = int(os.getenv("MAX_MOCK_FILES", "5"))
MOCK_LINE_LIMIT = int(os.getenv("MOCK_LINE_LIMIT", "100"))
MOCK_PATTERNS_PER_FILE = int(os.getenv("MOCK_PATTERNS_PER_FILE", "3"))
MOCK_OUTPUT_LIMIT = int(os.getenv("MOCK_OUTPUT_LIMIT", "20"))
PORT_RANGE_START = int(os.getenv("PORT_RANGE_START", "8000"))
PORT_RANGE_END = int(os.getenv("PORT_RANGE_END", "9000"))
DEFAULT_PORT = int(os.getenv("DEFAULT_PORT", "8000"))
GENERATED_FILES_DIR = os.getenv("GENERATED_FILES_DIR", "generated_files")

app = FastAPI(title="Impact Analysis Agent")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class AnalysisRequest(BaseModel):
    repo_url: str
    architecture_content: str
    prd_content: str = ""

class FileUploadResponse(BaseModel):
    success: bool
    extracted_text: str = ""
    filename: str = ""
    error: str = ""

class AnalysisResponse(BaseModel):
    success: bool
    analysis: str = ""
    document_id: str = ""
    timestamp: str = ""
    error: str = ""

class ImpactAnalysisAgent:
    def __init__(self):
        self.groq_api_key = GROQ_API_KEY
        self.model = GROQ_MODEL
        self.temperature = GROQ_TEMPERATURE
        self.groq_url = GROQ_URL

    def extract_text_from_file(self, file_data, filename):
        try:
            file_ext = os.path.splitext(filename)[1].lower()
            
            if file_ext == '.pdf' and PDF_SUPPORT:
                pdf_reader = PyPDF2.PdfReader(BytesIO(file_data))
                return "\n".join(page.extract_text() for page in pdf_reader.pages)
            elif file_ext in ['.doc', '.docx'] and DOC_SUPPORT:
                doc = Document(BytesIO(file_data))
                return "\n".join(paragraph.text for paragraph in doc.paragraphs)
            else:
                return file_data.decode('utf-8', errors='ignore')
        except Exception as e:
            raise Exception(f"Failed to extract text from {filename}: {str(e)}")

    def generate_pdf(self, content, title="Document"):
        if not PDF_SUPPORT:
            raise Exception("PDF generation not supported")
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = [Paragraph(title, styles['Title'])]
        
        for line in content.split('\n'):
            if line.strip():
                story.append(Paragraph(line, styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def detect_mock_data(self, content, file_path):
        """Detect mock data and hardcoded values in file content"""
        import re
        mock_patterns = [
            r'const\s+\w*[Mm]ock\w*\s*=',
            r'let\s+\w*[Mm]ock\w*\s*=',
            r'\w*[Dd]ummy\w*',
            r'\w*[Tt]est[Dd]ata\w*',
            r'localhost:\d+',
            r'127\.0\.0\.1',
            r'"http://.*"',
            r'\[\s*{[^}]*}\s*,\s*{[^}]*}\s*\]',  # Array of objects
            r'"\w+@example\.com"',
            r'"password"\s*:\s*"\w+"',
            r'api_key\s*[=:]\s*["\'][^"\']+["\']'
        ]
        
        found_mocks = []
        for i, line in enumerate(content.split('\n')[:MOCK_LINE_LIMIT], 1):
            for pattern in mock_patterns:
                if re.search(pattern, line, re.IGNORECASE):
                    found_mocks.append({
                        'line': i,
                        'content': line.strip()[:100],
                        'type': 'mock_data'
                    })
        
        return found_mocks if found_mocks else None

    def _format_mock_data(self, files):
        """Format mock data for prompt"""
        result = []
        for file_path, data in files.items():
            if data.get('mock_data'):
                result.append(f"\n{file_path}:")
                for mock in data['mock_data'][:MOCK_PATTERNS_PER_FILE]:
                    result.append(f"  Line {mock['line']}: {mock['content']}")
        return '\n'.join(result[:MOCK_OUTPUT_LIMIT])

    def clone_and_analyze_repo(self, repo_url):
        temp_dir = tempfile.mkdtemp()
        
        try:
            result = subprocess.run(['git', 'clone', '--depth', '1', repo_url, temp_dir], 
                                  check=True, capture_output=True, text=True, timeout=GIT_CLONE_TIMEOUT)
            
            repo_analysis = {'files': {}, 'structure': [], 'languages': set()}
            
            for root, dirs, files in os.walk(temp_dir):
                dirs[:] = [d for d in dirs if not d.startswith('.') and d not in ['node_modules', '__pycache__']]
                
                for file in files:
                    if file.startswith('.'):
                        continue
                        
                    file_path = os.path.join(root, file)
                    rel_path = os.path.relpath(file_path, temp_dir)
                    file_ext = os.path.splitext(file)[1].lower()
                    
                    repo_analysis['structure'].append(rel_path)
                    
                    if file_ext in ['.py', '.js', '.jsx', '.ts', '.tsx', '.java', '.go', '.rs', '.php', '.rb', '.cs', '.html', '.css']:
                        repo_analysis['languages'].add(file_ext)
                        
                        try:
                            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                                content = f.read()
                                repo_analysis['files'][rel_path] = {
                                    'content': content[:FILE_CONTENT_LIMIT],
                                    'full_size': len(content),
                                    'mock_data': self.detect_mock_data(content, rel_path)
                                }
                        except:
                            pass
            
            return repo_analysis
            
        except Exception as e:
            return {
                'files': {},
                'structure': ['Repository analysis unavailable'],
                'languages': {'.js', '.html', '.css'},
                'analysis_error': str(e)
            }
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    def call_groq_api(self, prompt):
        if not self.groq_api_key:
            raise Exception("GROQ_API_KEY not configured")
        
        # Clean the API key
        api_key = self.groq_api_key.strip().replace('\n', '').replace('\r', '')
        
        data = {
            "messages": [{"role": "user", "content": prompt}],
            "model": self.model,
            "temperature": self.temperature,
            "max_tokens": MAX_TOKENS
        }
        
        headers = {
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json',
            'User-Agent': 'Mozilla/5.0 (compatible)'
        }
        
        try:
            req = urllib.request.Request(
                self.groq_url,
                data=json.dumps(data).encode('utf-8'),
                headers=headers
            )
            
            with urllib.request.urlopen(req, timeout=API_REQUEST_TIMEOUT) as response:
                result = json.loads(response.read().decode('utf-8'))
                return result['choices'][0]['message']['content']
        except urllib.error.HTTPError as e:
            error_body = e.read().decode('utf-8')
            if e.code == 403:
                raise Exception(f"API key invalid or expired. Please check your Groq API key. Error: {error_body}")
            elif e.code == 429:
                raise Exception(f"Rate limit exceeded. Please try again later.")
            else:
                raise Exception(f"Groq API Error {e.code}: {error_body}")
        except Exception as e:
            raise Exception(f"API call failed: {str(e)}")

    def analyze_with_groq(self, repo_url, repo_analysis, architecture_content, prd_content=None):
        try:
            if len(architecture_content) > ARCH_CONTENT_LIMIT:
                architecture_content = architecture_content[:ARCH_CONTENT_LIMIT] + "\n[Content truncated]"
            
            key_files_content = ""
            for file_path, file_data in list(repo_analysis['files'].items())[:MAX_KEY_FILES]:
                key_files_content += f"\n=== {file_path} ===\n"
                key_files_content += file_data.get('content', '')[:KEY_FILE_CONTENT_LIMIT]
            
            languages = list(repo_analysis['languages'])[:MAX_LANGUAGES]
            
            # Extract mock data info
            mock_files = [f for f, data in repo_analysis['files'].items() if data.get('mock_data')][:MAX_MOCK_FILES]
            
            prompt = f"""Analyze repository for backend/database recommendations:

REPO: {repo_url}
LANGUAGES: {', '.join(languages)}
FILES ANALYZED: {len(repo_analysis['files'])}
MOCK DATA FILES: {len(mock_files)}

KEY FILES CONTENT:
{key_files_content}

MOCK DATA ANALYSIS:
{self._format_mock_data(repo_analysis['files'])}

ARCHITECTURE:
{architecture_content}

Provide:

# PROJECT SUMMARY
[Concise summary based on analyzed files]

# ARCHITECTURE DIAGRAM
[Create ASCII art architecture diagram showing:
- Frontend components
- Backend services
- Database layers
- API connections
- Data flow between components
Use boxes, arrows, and text to visualize the system architecture]

# BACKEND TECH STACK JUSTIFICATION
[Why specific backend technologies fit this project]

# DATABASE TECH STACK JUSTIFICATION
[Why specific database technologies fit this project]

# ALTERNATIVE BACKEND STACKS (3 options)
## Option 1: [Framework]
- Pros: [Specific advantages]
- Cons: [Limitations]
- Best for: [Use cases]

## Option 2: [Framework]
- Pros: [Specific advantages]
- Cons: [Limitations]
- Best for: [Use cases]

## Option 3: [Framework]
- Pros: [Specific advantages]
- Cons: [Limitations]
- Best for: [Use cases]

# ALTERNATIVE DATABASE STACKS (3 options)
## Option 1: [Database]
- Use cases: [Data patterns]
- Performance: [Characteristics]
- Best for: [Scenarios]

## Option 2: [Database]
- Use cases: [Data patterns]
- Performance: [Characteristics]
- Best for: [Scenarios]

## Option 3: [Database]
- Use cases: [Data patterns]
- Performance: [Characteristics]
- Best for: [Scenarios]

# DATABASE SCHEMA DESIGN
[Based on analyzed code patterns, design complete database schema:
- Tables/Collections with all fields and data types
- Primary keys, foreign keys, and relationships
- Indexes for performance optimization
- Constraints and validation rules]

# COMPLETE API ENDPOINTS
[All endpoints needed based on code analysis:
- Authentication endpoints (login, register, logout)
- CRUD endpoints for each data entity
- Business logic endpoints
- Integration endpoints
- Admin/management endpoints
Include HTTP methods, paths, request/response formats]

# IMPLEMENTATION GUIDE
## Phase 1: Setup
[Environment and project setup steps]

## Phase 2: Database
[Database implementation steps]

## Phase 3: Backend API
[API development steps with specific code examples]

## Phase 4: Integration
[Integration and testing steps]

Provide specific commands and code examples."""
            
            return self.call_groq_api(prompt)
            
        except Exception as e:
            return f"Error generating analysis: {str(e)}"

agent = ImpactAnalysisAgent()

@app.get("/favicon.ico")
async def favicon():
    return {"message": "No favicon"}

@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "version": "1.0",
        "groq_model": GROQ_MODEL,
        "api_key_configured": bool(GROQ_API_KEY and GROQ_API_KEY != "your_new_groq_api_key_here")
    }

@app.get("/", response_class=HTMLResponse)
async def read_root():
    html_content = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Impact Analysis Agent</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; background: #f5f5f5; }
        .container { max-width: 800px; margin: 0 auto; background: white; padding: 30px; border-radius: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }
        h1 { color: #333; text-align: center; margin-bottom: 30px; }
        .form-group { margin-bottom: 20px; }
        label { display: block; margin-bottom: 5px; font-weight: bold; color: #555; }
        input[type="url"], input[type="file"] { width: 100%; padding: 10px; border: 2px solid #ddd; border-radius: 5px; }
        button { background: #007bff; color: white; padding: 12px 30px; border: none; border-radius: 5px; cursor: pointer; font-size: 16px; }
        button:hover { background: #0056b3; }
        button:disabled { background: #ccc; cursor: not-allowed; }
        .status { margin-top: 20px; padding: 10px; border-radius: 5px; }
        .success { background: #d4edda; color: #155724; border: 1px solid #c3e6cb; }
        .error { background: #f8d7da; color: #721c24; border: 1px solid #f5c6cb; }
        .results { margin-top: 30px; padding: 20px; background: #f8f9fa; border-radius: 5px; }
        .btn-small { background: #28a745; color: white; padding: 5px 10px; margin: 2px; border: none; border-radius: 3px; cursor: pointer; }
    </style>
</head>
<body>
    <div class="container">
        <h1>🚀 Impact Analysis Agent</h1>
        <p style="text-align: center; color: #666; margin-bottom: 30px;">
            ✨ Intelligent GitHub Repository Analysis & Tech Stack Recommendations ✨
        </p>
        
        <form id="analysisForm">
            <div class="form-group">
                <label for="repoUrl">GitHub Repository URL:</label>
                <input type="url" id="repoUrl" placeholder="https://github.com/username/repository" required>
            </div>
            
            <div class="form-group">
                <label for="archFile">🏗️ Architecture Document (Required):</label>
                <input type="file" id="archFile" accept=".pdf,.doc,.docx,.txt,.json" required>
            </div>
            
            <div class="form-group">
                <label for="prdFile">📄 PRD Document (Optional):</label>
                <input type="file" id="prdFile" accept=".pdf,.doc,.docx,.txt,.json">
            </div>
            
            <button type="submit" id="analyzeBtn">🔍 Analyze Project</button>
        </form>
        
        <div id="status"></div>
        <div id="results" style="display: none;">
            <h3>📋 Analysis Results</h3>
            <div>
                <button class="btn-small" onclick="downloadDocument('prompt', 'txt')">Download TXT</button>
                <button class="btn-small" onclick="downloadDocument('prompt', 'json')">Download JSON</button>
                <button class="btn-small" onclick="downloadDocument('prompt', 'pdf')">Download PDF</button>
                <button class="btn-small" onclick="downloadDocument('prompt', 'docx')">Download DOCX</button>
            </div>
            <div id="analysisContent"></div>
        </div>
    </div>

    <script>
        let fileContents = { architecture: '', prd: '' };
        let currentDocumentId = null;

        document.getElementById('archFile').addEventListener('change', async (e) => {
            const file = e.target.files[0];
            if (file) await handleFileUpload(file, 'Architecture', 'architecture');
        });

        document.getElementById('prdFile').addEventListener('change', async (e) => {
            const file = e.target.files[0];
            if (file) await handleFileUpload(file, 'PRD', 'prd');
        });

        async function handleFileUpload(file, fileType, contentKey) {
            const status = document.getElementById('status');
            status.innerHTML = `<div style="color: #007bff;">📄 Processing ${fileType} file...</div>`;
            
            try {
                const formData = new FormData();
                formData.append('file', file);
                
                const response = await fetch('/upload-file', {
                    method: 'POST',
                    body: formData
                });
                
                const data = await response.json();
                
                if (data.success) {
                    fileContents[contentKey] = data.extracted_text;
                    status.innerHTML = `<div class="success">✅ ${fileType} file processed successfully!</div>`;
                } else {
                    throw new Error(data.error || 'File processing failed');
                }
            } catch (error) {
                fileContents[contentKey] = '';
                status.innerHTML = `<div class="error">❌ ${fileType} File Error: ${error.message}</div>`;
            }
        }

        document.getElementById('analysisForm').addEventListener('submit', async function(e) {
            e.preventDefault();
            
            const repoUrl = document.getElementById('repoUrl').value;
            const analyzeBtn = document.getElementById('analyzeBtn');
            const status = document.getElementById('status');
            const results = document.getElementById('results');
            
            if (!fileContents.architecture) {
                status.innerHTML = '<div class="error">❌ Please upload an Architecture document first!</div>';
                return;
            }
            
            analyzeBtn.disabled = true;
            analyzeBtn.textContent = '🔄 Analyzing...';
            status.innerHTML = '<div style="color: #007bff;">🔄 Analyzing repository...</div>';
            
            try {
                const response = await fetch('/analyze', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({
                        repo_url: repoUrl,
                        architecture_content: fileContents.architecture,
                        prd_content: fileContents.prd
                    })
                });
                
                const data = await response.json();
                
                if (data.success) {
                    currentDocumentId = data.document_id;
                    document.getElementById('analysisContent').innerHTML = 
                        `<pre style="background: #f8f9fa; padding: 15px; border-radius: 5px; white-space: pre-wrap;">${data.analysis}</pre>`;
                    results.style.display = 'block';
                    status.innerHTML = '<div class="success">✅ Analysis completed successfully!</div>';
                } else {
                    throw new Error(data.error || 'Analysis failed');
                }
            } catch (error) {
                status.innerHTML = `<div class="error">❌ Error: ${error.message}</div>`;
            } finally {
                analyzeBtn.disabled = false;
                analyzeBtn.textContent = '🔍 Analyze Project';
            }
        });

        async function downloadDocument(docType, format) {
            if (!currentDocumentId) {
                alert('No documents available for download');
                return;
            }
            
            try {
                const url = `/download/${currentDocumentId}/${docType}/${format}`;
                const response = await fetch(url);
                
                if (!response.ok) throw new Error(`Download failed: ${response.status}`);
                
                const blob = await response.blob();
                const downloadUrl = window.URL.createObjectURL(blob);
                
                const link = document.createElement('a');
                link.href = downloadUrl;
                link.download = `${docType}_${currentDocumentId}.${format}`;
                document.body.appendChild(link);
                link.click();
                document.body.removeChild(link);
                
                window.URL.revokeObjectURL(downloadUrl);
            } catch (error) {
                alert(`Download failed: ${error.message}`);
            }
        }
    </script>
</body>
</html>
    """
    return HTMLResponse(content=html_content)

@app.post("/upload-file", response_model=FileUploadResponse)
async def upload_file(file: UploadFile = File(...)):
    try:
        file_data = await file.read()
        extracted_text = agent.extract_text_from_file(file_data, file.filename)
        
        return FileUploadResponse(
            success=True,
            extracted_text=extracted_text,
            filename=file.filename
        )
    except Exception as e:
        return FileUploadResponse(
            success=False,
            error=str(e)
        )

@app.post("/analyze", response_model=AnalysisResponse)
async def analyze_project(request: AnalysisRequest):
    try:
        # Validate repository URL format
        if not request.repo_url or not request.repo_url.startswith(('https://github.com/', 'http://github.com/')):
            raise HTTPException(status_code=400, detail="Invalid GitHub repository URL format")
        
        if not request.architecture_content or len(request.architecture_content.strip()) < 10:
            raise HTTPException(status_code=400, detail="Architecture document content is required and must be substantial")
        
        # Check API key configuration
        if not GROQ_API_KEY or GROQ_API_KEY == "your_new_groq_api_key_here":
            raise HTTPException(status_code=500, detail="Groq API key not properly configured")
        
        repo_analysis = agent.clone_and_analyze_repo(request.repo_url)
        analysis = agent.analyze_with_groq(
            request.repo_url, 
            repo_analysis, 
            request.architecture_content, 
            request.prd_content
        )
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        os.makedirs(GENERATED_FILES_DIR, exist_ok=True)
        
        documents = {
            'prompt': analysis,
            'architecture': request.architecture_content,
            'prd': request.prd_content or 'No PRD document provided',
            'repository_url': request.repo_url,
            'analysis_metadata': {
                'timestamp': datetime.now().isoformat(),
                'files_analyzed': len(repo_analysis.get('files', {})),
                'languages_detected': list(repo_analysis.get('languages', [])),
                'mock_data_files': len([f for f, data in repo_analysis.get('files', {}).items() if data.get('mock_data')])
            }
        }
        
        with open(f'{GENERATED_FILES_DIR}/documents_{timestamp}.json', 'w', encoding='utf-8') as f:
            json.dump(documents, f, indent=2)
        
        return AnalysisResponse(
            success=True,
            analysis=analysis,
            document_id=timestamp,
            timestamp=datetime.now().isoformat()
        )
        
    except HTTPException:
        raise
    except Exception as e:
        return AnalysisResponse(
            success=False,
            error=str(e)
        )

@app.get("/download/{document_id}/{doc_type}/{format_type}")
async def download_document(document_id: str, doc_type: str, format_type: str):
    try:
        json_path = f'{GENERATED_FILES_DIR}/documents_{document_id}.json'
        
        if not os.path.exists(json_path):
            raise HTTPException(status_code=404, detail="Document not found")
        
        with open(json_path, 'r', encoding='utf-8') as f:
            documents = json.load(f)
        
        if doc_type not in documents:
            raise HTTPException(status_code=404, detail="Document type not found")
        
        content = documents[doc_type]
        filename = f"{doc_type}_{document_id}"
        
        if format_type == 'json':
            json_content = json.dumps({"content": content}, indent=2)
            return StreamingResponse(
                BytesIO(json_content.encode('utf-8')),
                media_type='application/json',
                headers={"Content-Disposition": f"attachment; filename={filename}.json"}
            )
        elif format_type == 'txt':
            return StreamingResponse(
                BytesIO(content.encode('utf-8')),
                media_type='text/plain',
                headers={"Content-Disposition": f"attachment; filename={filename}.txt"}
            )
        elif format_type == 'pdf':
            if not PDF_SUPPORT:
                raise HTTPException(status_code=400, detail="PDF generation not supported")
            pdf_data = agent.generate_pdf(content, f"{doc_type.title()} - {document_id}")
            return StreamingResponse(
                BytesIO(pdf_data),
                media_type='application/pdf',
                headers={"Content-Disposition": f"attachment; filename={filename}.pdf"}
            )
        elif format_type == 'docx':
            if not DOC_SUPPORT:
                raise HTTPException(status_code=400, detail="DOCX generation not supported")
            docx_data = agent.generate_docx(content, f"{doc_type.title()} - {document_id}")
            return StreamingResponse(
                BytesIO(docx_data),
                media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={"Content-Disposition": f"attachment; filename={filename}.docx"}
            )
        else:
            raise HTTPException(status_code=400, detail="Invalid format")
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import socket
    
    # Find available port
    def find_free_port():
        for port in range(PORT_RANGE_START, PORT_RANGE_END):
            try:
                with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                    s.bind(('127.0.0.1', port))
                    return port
            except OSError:
                continue
        return DEFAULT_PORT
    
    port = int(os.environ.get("PORT", find_free_port()))
    # Use 127.0.0.1 for local development, 0.0.0.0 for production
    host = "127.0.0.1" if os.environ.get("PORT") is None else "0.0.0.0"
    print(f"Starting server on http://{host}:{port}")
    uvicorn.run(app, host=host, port=port)